"""Utility functions for processing uploaded files and extracting text content."""
import io
import os
from typing import List, Dict, Optional

def extract_text_from_file(file_data: bytes, file_name: str, file_type: str) -> str:
    """
    Extract text content from various file types.
    
    Args:
        file_data: Binary file data
        file_name: Name of the file
        file_type: MIME type of the file
        
    Returns:
        Extracted text content
    """
    file_ext = os.path.splitext(file_name)[1].lower()
    
    try:
        # Text files
        if file_ext in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
            try:
                # Try UTF-8 first
                return file_data.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                return file_data.decode('latin-1', errors='ignore')
        
        # PDF files
        elif file_ext == '.pdf':
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                return f"[PDF file: {file_name} - PyPDF2 not installed. Install with: pip install PyPDF2]"
            except Exception as e:
                return f"[Error reading PDF {file_name}: {str(e)}]"
        
        # Word documents
        elif file_ext in ['.docx', '.doc']:
            try:
                if file_ext == '.docx':
                    from docx import Document
                    doc = Document(io.BytesIO(file_data))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            text += "\n" + " | ".join([cell.text for cell in row.cells])
                    return text
                else:
                    # .doc files require python-docx2txt or antiword
                    return f"[DOC file: {file_name} - DOCX format preferred. Please convert to DOCX.]"
            except ImportError:
                return f"[Word document: {file_name} - python-docx not installed. Install with: pip install python-docx]"
            except Exception as e:
                return f"[Error reading Word document {file_name}: {str(e)}]"
        
        # Excel files
        elif file_ext in ['.xlsx', '.xls']:
            try:
                import pandas as pd
                excel_file = io.BytesIO(file_data)
                # Read all sheets
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
                text = ""
                for sheet_name, df in excel_data.items():
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    text += df.to_string(index=False) + "\n"
                return text
            except ImportError:
                return f"[Excel file: {file_name} - openpyxl/pandas not installed. Install with: pip install openpyxl pandas]"
            except Exception as e:
                return f"[Error reading Excel file {file_name}: {str(e)}]"
        
        # Images (OCR would require additional setup)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            return f"[Image file: {file_name} - Image content not extracted. OCR not implemented.]"
        
        # Default: try to decode as text
        else:
            try:
                return file_data.decode('utf-8', errors='ignore')
            except:
                return f"[Binary file: {file_name} - Unable to extract text content]"
    
    except Exception as e:
        return f"[Error processing file {file_name}: {str(e)}]"


def process_uploaded_files(uploaded_files: List) -> Dict[str, str]:
    """
    Process a list of uploaded files and extract text content.
    
    Args:
        uploaded_files: List of uploaded file objects (from Streamlit)
        
    Returns:
        Dictionary mapping file names to extracted text content
    """
    file_contents = {}
    
    for file in uploaded_files:
        try:
            file_data = file.getvalue()
            file_name = file.name
            file_type = file.type if hasattr(file, 'type') else 'application/octet-stream'
            
            text_content = extract_text_from_file(file_data, file_name, file_type)
            file_contents[file_name] = text_content
            
        except Exception as e:
            file_contents[file.name] = f"[Error processing {file.name}: {str(e)}]"
    
    return file_contents


def format_file_content_for_llm(file_contents: Dict[str, str], max_length: int = 8000) -> str:
    """
    Format file contents for inclusion in LLM query.
    
    Args:
        file_contents: Dictionary mapping file names to text content
        max_length: Maximum total length of formatted content
        
    Returns:
        Formatted string for LLM context
    """
    if not file_contents:
        return ""
    
    formatted_parts = []
    total_length = 0
    
    for file_name, content in file_contents.items():
        # Truncate individual file content if too long
        if len(content) > max_length // len(file_contents):
            content = content[:max_length // len(file_contents)] + "... [truncated]"
        
        file_summary = f"--- File: {file_name} ---\n{content}\n"
        
        if total_length + len(file_summary) > max_length:
            # Add truncated indicator
            remaining = max_length - total_length
            if remaining > 50:
                file_summary = file_summary[:remaining] + "... [additional content truncated]"
            else:
                break
        
        formatted_parts.append(file_summary)
        total_length += len(file_summary)
    
    return "\n".join(formatted_parts)

